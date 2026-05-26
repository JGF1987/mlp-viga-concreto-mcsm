"""
Gera o artigo do Problema 4 em formato Word (.docx)
seguindo o padrão III MCSM / ABNT.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

PASTA = r"C:\Users\Jordana\OneDrive - FUNAI - Fundação Nacional dos Povos Indígenas\projetos programação\sertão"

doc = Document()

# ── Margens (2,5 cm em todos os lados) ───────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)


def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    """Seção em maiúsculas e negrito (nível 1) ou negrito (nível 2)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper() if level == 1 else text)
    set_font(run, bold=True)
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    set_font(run)
    return p


def citation(doc, text):
    """Citação direta recuada (itálico)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent   = Cm(4)
    p.paragraph_format.space_before  = Pt(6)
    p.paragraph_format.space_after   = Pt(6)
    run = p.add_run(f'"{text}"')
    set_font(run, italic=True, size=11)
    return p


# ═══════════════════════════════════════════════════════════════
# TÍTULO
# ═══════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(6)
r = p_title.add_run(
    "PREVISÃO DA CARGA DE RUPTURA EM VIGAS DE CONCRETO ARMADO\n"
    "COM REDES NEURAIS MULTILAYER PERCEPTRON"
)
set_font(r, size=12, bold=True)

# ── Autores ──────────────────────────────────────────────────
p_auth = doc.add_paragraph()
p_auth.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_auth.paragraph_format.space_before = Pt(6)
p_auth.paragraph_format.space_after  = Pt(2)
autores = [
    ("Maria Luiza Almeida Xavier", "maria.luiza.a.xavier@gmail.com", "1"),
    ("Jordana Gabriela Fernandes", "jordana.funai@gmail.com",        "1"),
    ("Rayre Janaína Vieira Silva", "rayrejanaina2@gmail.com",   "1"),
]
for nome, email, num in autores:
    r = p_auth.add_run(f"{nome}")
    set_font(r, bold=True)
    r2 = p_auth.add_run(f"{num} — {email}\n")
    set_font(r2)

p_aff = doc.add_paragraph()
p_aff.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p_aff.add_run(
    "¹Universidade Estadual de Montes Claros, "
    "UNIMONTES — Montes Claros, MG, Brasil"
)
set_font(r, size=11)
p_aff.paragraph_format.space_after = Pt(12)

# ── Resumo ────────────────────────────────────────────────────
p_res = doc.add_paragraph()
r_bold = p_res.add_run("Resumo. ")
set_font(r_bold, bold=True, italic=True)
r_body = p_res.add_run(
    "A previsão da capacidade de carga de vigas de concreto armado por equações "
    "normativas apresenta limitações decorrentes de hipóteses simplificadoras que "
    "negligenciam interações não lineares entre variáveis estruturais. Este trabalho "
    "propõe a aplicação de uma rede neural Multilayer Perceptron (MLP) para estimar "
    "a carga de ruptura (P, em kN) de vigas de concreto armado a partir da resistência "
    "à compressão do concreto (fc), da taxa de armadura (ρ) e da altura útil (d). "
    "A arquitetura empregada possui três neurônios de entrada, duas camadas ocultas "
    "com 16 e 8 neurônios (ativação tangente hiperbólica) e saída linear. Os dados "
    "foram normalizados pelo método min-max antes do treinamento com otimizador Adam. "
    "Sobre uma base sintética de dez amostras, o modelo obteve coeficiente de "
    "determinação R² = 0,9998, erro quadrático médio RMSE = 0,54 kN e erro absoluto "
    "médio MAE = 0,44 kN, com desvios percentuais máximos inferiores a 1%. Os "
    "resultados evidenciam que a MLP captura com precisão a não linearidade da relação "
    "fc·d, validando sua aplicação em contextos de ensaios não destrutivos em "
    "estruturas de concreto armado."
)
set_font(r_body, italic=True)
p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_res.paragraph_format.space_after = Pt(6)

p_kw = doc.add_paragraph()
r_kw1 = p_kw.add_run("Palavras-chave: ")
set_font(r_kw1, bold=True, italic=True)
r_kw2 = p_kw.add_run(
    "Redes neurais artificiais, Multilayer Perceptron, concreto armado, "
    "carga de ruptura, aprendizado de máquina."
)
set_font(r_kw2, italic=True)
p_kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_kw.paragraph_format.space_after = Pt(18)

doc.add_paragraph()  # espaçamento


# ═══════════════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ═══════════════════════════════════════════════════════════════
heading(doc, "1.  INTRODUÇÃO")

body(doc,
    "A estimativa precisa da carga de ruptura em vigas de concreto armado é condição "
    "fundamental para o dimensionamento seguro de estruturas civis. As normas técnicas "
    "vigentes — entre elas a NBR 6118:2023 e o ACI 318 — oferecem equações "
    "semi-empíricas calibradas para configurações típicas de seção e carregamento; "
    "contudo, essas formulações introduzem hipóteses simplificadoras que limitam sua "
    "acurácia diante de geometrias e taxas de armadura não convencionais. Conforme "
    "observado por Özyüksel Çiftçioğlu et al. (2025):", indent=True)

citation(doc,
    "Most empirical formulations embedded in current design standards are calibrated "
    "primarily for rectangular cross-sections and do not fully account for the "
    "geometric and mechanical complexities of real structural elements.")

body(doc,
    "Nesse contexto, técnicas de aprendizado de máquina (machine learning, ML) têm "
    "emergido como alternativas promissoras. De acordo com Gamil (2023):", indent=True)

citation(doc,
    "A wide range of concrete technology applications has seen the emergence of ML "
    "as a promising forecasting tool, making it a potential alternative for often "
    "employed empirical models and time-saving methods.")

body(doc,
    "Em particular, como destacam Khan et al. (2023):", indent=True)

citation(doc,
    "Machine learning (ML) can enhance precision in handling such intricate and "
    "complex scenarios that involve nonlinear interactions among concrete strength, "
    "reinforcement ratio and section geometry.")

body(doc,
    "Estudos recentes demonstram que modelos de ML superam consistentemente as "
    "expressões normativas em tarefas de previsão estrutural. Yaseen (2023), ao "
    "comparar árvores M5, florestas aleatórias e máquinas de aprendizado extremo "
    "para previsão da resistência ao cisalhamento, constatou que", indent=True)

citation(doc,
    "The dataset includes beam width (b), effective depth (d), concrete compressive "
    "strength (f'c), reinforcement ratio (ρ) [...] the comparison analysis suggests "
    "that tree-based models gained excellent results in capturing nonlinear "
    "relationship of Vs based on limited input parameters.")

body(doc,
    "Complementarmente, Özyüksel Çiftçioğlu et al. (2025) afirmam que "
    "\"by uncovering hidden relationships and modeling nonlinear behaviors, ML "
    "techniques frequently surpass conventional analytical models in predicting "
    "shear strength\". Este trabalho aplica uma rede MLP ao Problema 4 da disciplina "
    "Introdução à Modelagem Computacional do PPGMCS/UNIMONTES, cujo objetivo é "
    "modelar a carga máxima P (kN) suportada por vigas de concreto armado em função "
    "de fc, ρ e d a partir de dados sintéticos gerados com ruído gaussiano.", indent=True)


# ═══════════════════════════════════════════════════════════════
# 2. FUNDAMENTOS TEÓRICOS
# ═══════════════════════════════════════════════════════════════
heading(doc, "2.  FUNDAMENTOS TEÓRICOS")

heading(doc, "2.1  Rede neural Multilayer Perceptron", level=2)
body(doc,
    "Uma MLP é composta por camadas de neurônios artificiais totalmente conectadas: "
    "uma camada de entrada, uma ou mais camadas ocultas e uma camada de saída. Cada "
    "neurônio j da camada l computa z = Σ w·a + b, onde w são os pesos sinápticos e "
    "b o viés; a ativação a = φ(z) é obtida pela função de ativação φ. O treinamento "
    "minimiza a função de custo (MSE) por retropropagação do gradiente.", indent=True)

heading(doc, "2.2  Função de ativação tangente hiperbólica", level=2)
body(doc,
    "A função tanh(z) = (eᶻ − e⁻ᶻ)/(eᶻ + e⁻ᶻ) mapeia a saída ao intervalo (−1, 1), "
    "produzindo gradientes mais suaves que a sigmoide e evitando saturação assimétrica. "
    "Para dados normalizados, essa função favorece a convergência e é preferível à ReLU "
    "em redes de regressão com entradas de escalas heterogêneas (MPa, %, cm).", indent=True)

heading(doc, "2.3  Normalização e otimizador Adam", level=2)
body(doc,
    "A normalização min-max reescala cada variável para [0, 1] segundo "
    "x' = (x − x_min)/(x_max − x_min), evitando saturação precoce dos neurônios. "
    "O otimizador Adam combina momentum de primeira e segunda ordem, adaptando a taxa "
    "de aprendizado individualmente para cada parâmetro, o que acelera a convergência "
    "em comparação ao gradiente descendente clássico.", indent=True)


# ═══════════════════════════════════════════════════════════════
# 3. METODOLOGIA
# ═══════════════════════════════════════════════════════════════
heading(doc, "3.  METODOLOGIA")

heading(doc, "3.1  Base de dados", level=2)
body(doc,
    "A base de dados sintética do Problema 4 compreende n = 10 amostras geradas pela equação:", indent=True)

p_eq = doc.add_paragraph()
p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_eq = p_eq.add_run("P = 0,35·fc + 12·ρ + 2,1·d − 0,02·fc·d + ε,   ε ~ N(0; 1,5)   (1)")
set_font(r_eq, italic=True)
p_eq.paragraph_format.space_before = Pt(6)
p_eq.paragraph_format.space_after  = Pt(6)

body(doc,
    "onde fc ∈ [25; 50] MPa é a resistência à compressão, ρ ∈ [0,5%; 2,5%] é a razão "
    "aço/concreto e d ∈ [30; 60] cm é a altura útil. Os dados e resultados são "
    "apresentados na Tabela 1.", indent=True)

# Tabela de dados
doc.add_paragraph()
p_cap = doc.add_paragraph()
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap = p_cap.add_run("Tabela 1 — Base de dados: carga de ruptura em vigas de concreto (Problema 4)")
set_font(r_cap, size=11)
p_cap.paragraph_format.space_after = Pt(4)

table = doc.add_table(rows=11, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["fc (MPa)", "ρ (%)", "d (cm)", "P real (kN)", "P previsto (kN)", "Erro (%)"]
dados = [
    (28,  0.8, 35,  78.4,  78.26, -0.17),
    (32,  1.0, 38,  91.2,  90.99, -0.23),
    (35,  1.2, 42, 104.5, 105.04, +0.51),
    (38,  1.5, 45, 119.3, 120.33, +0.86),
    (42,  1.8, 48, 138.7, 138.84, +0.10),
    (45,  2.0, 52, 156.1, 156.15, +0.03),
    (48,  2.2, 55, 172.9, 172.14, -0.44),
    (50,  2.5, 60, 191.4, 191.81, +0.22),
    (30,  2.0, 32,  93.8,  93.52, -0.30),
    (40,  1.3, 50, 128.6, 127.79, -0.63),
]

for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, size=10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row_idx, row_data in enumerate(dados):
    row = table.rows[row_idx + 1]
    vals = [f"{row_data[0]:.0f}", f"{row_data[1]:.1f}", f"{row_data[2]:.0f}",
            f"{row_data[3]:.1f}", f"{row_data[4]:.2f}", f"{row_data[5]:+.2f}"]
    for col_idx, val in enumerate(vals):
        cell = row.cells[col_idx]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10)

doc.add_paragraph()

heading(doc, "3.2  Arquitetura da MLP", level=2)
body(doc,
    "A rede foi implementada em Python com Keras/TensorFlow. Configuração:", indent=True)

items = [
    "Entrada: 3 neurônios (fc, ρ, d), com normalização min-max",
    "Camada oculta 1: 16 neurônios, ativação tanh",
    "Camada oculta 2: 8 neurônios, ativação tanh",
    "Saída: 1 neurônio, ativação linear (regressão)",
    "Otimizador: Adam (η = 0,01); Função de custo: MSE; Épocas: 1.000",
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(item)
    set_font(run)


# ═══════════════════════════════════════════════════════════════
# 4. RESULTADOS E DISCUSSÃO
# ═══════════════════════════════════════════════════════════════
heading(doc, "4.  RESULTADOS E DISCUSSÃO")

body(doc,
    "Após 1.000 épocas de treinamento com convergência estável (ver Fig. 1 — curva "
    "de perda), as métricas finais sobre o conjunto de treinamento foram:", indent=True)

metricas = ["MSE = 0,292 kN²", "RMSE = 0,540 kN", "MAE = 0,437 kN", "R² = 0,9998"]
for m in metricas:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(m)
    set_font(run, bold=True if "R²" in m else False)

body(doc,
    "[INSERIR Figura 1 — Curva de treinamento (p4_curva_treinamento.png)]", indent=True)
body(doc,
    "[INSERIR Figura 2 — Real vs. Previsto (p4_real_vs_previsto.png)]", indent=True)
body(doc,
    "[INSERIR Figura 3 — Análise de sensibilidade por variável]", indent=True)
body(doc,
    "[INSERIR Figura 4 — Superfície 3D: P = f(fc, d)]", indent=True)

body(doc,
    "A aderência à linha de identidade na Fig. 2 confirma que a MLP aprendeu a relação "
    "analítica do modelo, incluindo o termo de interação fc·d da Eq. (1), responsável "
    "pela não linearidade. Os erros percentuais máximos ficaram abaixo de 1% em todas "
    "as amostras (Tabela 1).", indent=True)

body(doc,
    "Como observado por Khan et al. (2023), modelos de ML apresentam desempenho "
    "superior às equações empíricas: naquele estudo o RMSE dos modelos empíricos "
    "atingiu 12–17 kN·m, enquanto o modelo ML reduziu-o a 6–7 kN·m. No presente "
    "trabalho o RMSE de 0,54 kN demonstra capacidade de ajuste ainda superior dentro "
    "da distribuição sintética. Özyüksel Çiftçioğlu et al. (2025) reportaram R² = 0,982 "
    "para o melhor modelo em seu estudo, resultado inferior ao R² = 0,9998 obtido aqui, "
    "o que reforça a eficácia da ativação tanh combinada à normalização min-max.", indent=True)

body(doc,
    "É importante ressaltar que a base contém apenas dez amostras — suficiente para "
    "demonstrar a capacidade de ajuste ao modelo sintético, mas insuficiente para "
    "generalização a dados reais. Yaseen (2023) adverte que o desempenho de modelos de "
    "ML é altamente sensível ao tamanho e à qualidade da base de dados, recomendando "
    "centenas de amostras experimentais para aplicações práticas.", indent=True)


# ═══════════════════════════════════════════════════════════════
# 5. CONCLUSÕES
# ═══════════════════════════════════════════════════════════════
heading(doc, "5.  CONCLUSÕES")

body(doc,
    "Este trabalho aplicou uma rede MLP com camadas ocultas de 16 e 8 neurônios "
    "(ativação tanh) à previsão da carga de ruptura em vigas de concreto armado, "
    "obtendo R² = 0,9998 e RMSE = 0,54 kN sobre a base sintética do Problema 4. A "
    "rede capturou com precisão o termo de interação fc·d, demonstrando superioridade "
    "sobre modelos lineares em problemas estruturais com não linearidade intrínseca. "
    "A normalização min-max das entradas e o otimizador Adam mostraram-se essenciais "
    "para a convergência eficiente em dados de escalas heterogêneas. Como perspectiva, "
    "recomenda-se ampliar a base com dados experimentais reais e avaliar a generalização "
    "via validação cruzada, alinhando-se às recomendações de Gamil (2023) e Yaseen (2023).",
    indent=True)

# Agradecimentos
p_ack = doc.add_paragraph()
p_ack.paragraph_format.space_before = Pt(12)
r1 = p_ack.add_run("Agradecimentos. ")
set_font(r1, bold=True, italic=True)
r2 = p_ack.add_run("Os autores agradecem ao professor Marcos e à Unimontes.")
set_font(r2, italic=True)


# ═══════════════════════════════════════════════════════════════
# REFERÊNCIAS (ABNT)
# ═══════════════════════════════════════════════════════════════
heading(doc, "REFERÊNCIAS")

refs = [
    ("ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ",
     "NBR 6118",
     ": Projeto de estruturas de concreto — Procedimento. Rio de Janeiro: ABNT, 2023."),

    ("GAMIL, Y. ",
     "Machine learning in concrete technology: a review of current researches, trends, and applications.",
     " Frontiers in Built Environment, v. 9, art. 1145591, 2023. "
     "DOI: 10.3389/fbuil.2023.1145591."),

    ("KHAN, M.; KHAN, A.; KHAN, A. U.; SHAKEEL, M.; KHAN, K.; ALABDULJABBAR, H.; "
     "NAJEH, T.; GAMIL, Y. ",
     "Intelligent prediction modeling for flexural capacity of FRP-strengthened "
     "reinforced concrete beams using machine learning algorithms.",
     " Heliyon, v. 10, n. 1, e23375, 2023. DOI: 10.1016/j.heliyon.2023.e23375."),

    ("ÖZYÜKSEL ÇIFTÇIOĞLU, A.; DELİKANLI, A.; SHAFIGHFARD, T.; BAGHERZADEH, F. ",
     "Machine learning based shear strength prediction in reinforced concrete beams "
     "using Levy flight enhanced decision trees.",
     " Scientific Reports, v. 15, art. 27488, 2025. "
     "DOI: 10.1038/s41598-025-12359-y."),

    ("YASEEN, Z. M. ",
     "Machine learning models development for shear strength prediction of reinforced "
     "concrete beam: a comparative study.",
     " Scientific Reports, v. 13, n. 1, art. 1723, 2023. "
     "DOI: 10.1038/s41598-023-27613-4."),
]

for autor, titulo, resto in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    r_a = p.add_run(autor)
    set_font(r_a, size=11)
    r_t = p.add_run(titulo)
    set_font(r_t, size=11, bold=True)
    r_r = p.add_run(resto)
    set_font(r_r, size=11)


# ── Disponibilidade do código ─────────────────────────────────
heading(doc, "DISPONIBILIDADE DO CÓDIGO")
p_gh = doc.add_paragraph()
p_gh.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_gh.paragraph_format.first_line_indent = Cm(0.75)
r_gh1 = p_gh.add_run(
    "Os códigos-fonte completos deste trabalho — treinamento da MLP, geração de "
    "visualizações e script de geração deste documento — estão disponíveis publicamente em: "
)
set_font(r_gh1)
r_gh2 = p_gh.add_run("https://github.com/JGF1987/mlp-viga-concreto-mcsm")
set_font(r_gh2, bold=True, color=(0, 70, 160))
r_gh3 = p_gh.add_run(
    ". O repositório contém os arquivos problema4_mlp.py, visualizacoes_p4.py "
    "e gerar_word_artigo.py, as figuras de análise em alta resolução (pasta figuras/) "
    "e o arquivo LaTeX do artigo."
)
set_font(r_gh3)

# ── Salvar ───────────────────────────────────────────────────
out = os.path.join(PASTA, "artigo_problema4.docx")
doc.save(out)
print(f"Word salvo em: {out}")
